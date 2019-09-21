import random
import string
import time
import docx

def getReportName(randomlength=16):
    """
    生成一个指定长度的随机字符串，其中
    string.digits=0123456789
    string.ascii_letters=abcdefghigklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ
    """
    str_list = [random.choice(string.digits + string.ascii_letters) for i in range(randomlength)]
    random_str = ''.join(str_list)
    return random_str

def make(context):
    content = context['text']
    content_str = "检测文章内容：" + str(content)
    copytitle = context['name']
    copytitle_str = "相似文章标题：" + str(copytitle)
    copynum = context['num'] * 100
    copynum_str = "复写率：" + str(copynum) + "%"
    selfnum = (100 - float(copynum))
    selfnum_str = "自写率：" + str(selfnum) + "%"
    reportname = getReportName()
    reportname_str = "报告编号：" + str(reportname)
    # 格式化成2016-03-20 11:45:39形式
    reporttime = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
    reporttime_str = "检测时间：" + str(reporttime)
    reportlen = len(content)
    reportlen_str = "检测字数：" + str(reportlen)

    file = docx.Document()
    file.add_paragraph(reportname_str)
    file.add_paragraph(reporttime_str)
    file.add_paragraph(reportlen_str)
    file.add_paragraph('-----------------------------------')
    file.add_paragraph(content_str)
    file.add_paragraph('-----------------------------------')
    file.add_paragraph(copytitle_str)
    file.add_paragraph("相似比超过20%的片段：")
    for p in context['par']:
        a = str(p[1]) + "    ---- 相似度：" + str(p[0])
        file.add_paragraph(a)
        #print(p)
    file.add_paragraph('-----------------------------------')
    file.add_paragraph(copynum_str)
    file.add_paragraph(selfnum_str)

    doc_name = reportname+".docx"
    file.save(doc_name)

    print(reportname_str)
    print(reporttime_str)
    print(reportlen_str)
    print('-----------------------------------')
    print(content_str)
    print('-----------------------------------')
    print(copytitle_str)
    print(copynum_str)
    print(selfnum_str)

